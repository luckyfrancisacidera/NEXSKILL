"""Text extraction and normalization helpers shared by the resume parser pipeline."""

import io
import re
from typing import List, Optional, Union, Tuple
import pdfplumber
from docx import Document
from itertools import groupby


# =========================================
# TEXT NORMALIZATION PRIMITIVES
# =========================================
RE_HYPHEN = re.compile(r"(\w)-\s*\n\s*(\w)")
RE_SPACES = re.compile(r"[ \t]+")
RE_NEWLINES = re.compile(r"\n{3,}")


def _words_to_lines(words: List[dict], y_tol: float = 3.0) -> List[str]:
    """Reconstruct lines from PDF word fragments when direct text extraction is sparse."""
    if not words:
        return []

    words.sort(key=lambda w: (w.get("top", 0.0), w.get("x0", 0.0)))
    out = []

    for _, group in groupby(words, key=lambda w: round(w.get("top", 0.0) / y_tol)):
        line = " ".join((w.get("text") or "").strip() for w in group).strip()
        if line:
            out.append(line)

    return out


def extract_text_from_upload(filename: str, content: bytes) -> str:
    """Normalize file ingestion across PDF, DOCX, and plain-text uploads."""
    name = filename.lower()
    buf = io.StringIO()

    if name.endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = (page.extract_text() or "").strip()
                if not text:
                    # PDF extraction can fail on layout-heavy resumes, so we fall back
                    # to grouped word lines to preserve enough structure for parsing.
                    words = page.extract_words(use_text_flow=True) or []
                    text = "\n".join(_words_to_lines(words)).strip()
                if text:
                    buf.write(text + "\n")
        return buf.getvalue().strip()

    if name.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        buf.write("\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip()))
        for table in doc.tables:
            for row in table.rows:
                buf.write(
                    "\n".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                )
        return buf.getvalue().strip()

    return content.decode("utf-8", errors="ignore")


def normalize_text(text: str) -> str:
    """Collapse noisy whitespace and line-break artifacts before section parsing."""
    text = (text or "").replace("\x00", " ")
    text = RE_HYPHEN.sub(r"\1\2", text)
    text = RE_SPACES.sub(" ", text)
    text = RE_NEWLINES.sub("\n\n", text)
    return text.strip()


# =========================================
# YEARS OF EXPERIENCE HELPERS
# =========================================

WORD_NUMBERS = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
    "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
    "eleven": 11, "twelve": 12, "thirteen": 13, "fourteen": 14,
    "fifteen": 15, "sixteen": 16, "seventeen": 17,
    "eighteen": 18, "nineteen": 19, "twenty": 20
}

RE_YEARS_EXPERIENCE = re.compile(
    r"""
    (?:
        (?P<num>\d{1,2})\s*(?:\+|\bto\b|-)?\s*(?P<num2>\d{1,2})?
        |
        (?P<word>one|two|three|four|five|six|seven|eight|nine|ten|
                 eleven|twelve|thirteen|fourteen|fifteen|sixteen|
                 seventeen|eighteen|nineteen|twenty)
    )
    \s*
    (?:\+?\s*)?
    (?:years?|yrs?)
    \s+
    (?:of\s+)?
    experience
    """,
    re.IGNORECASE | re.VERBOSE
)


def extract_years_of_experience(
    text: str
) -> Optional[Union[int, Tuple[int, int]]]:
    """Extract explicit years-of-experience requirements from freeform text."""
    matches = []

    for m in RE_YEARS_EXPERIENCE.finditer(text):
        if m.group("num"):
            start = int(m.group("num"))
            end = m.group("num2")
            matches.append((start, int(end)) if end else start)

        elif m.group("word"):
            matches.append(WORD_NUMBERS[m.group("word").lower()])

    if not matches:
        return None

    # Prefer the highest requirement because job descriptions often repeat
    # experience in multiple forms and the strictest mention should win.
    def max_value(v):
        return v[1] if isinstance(v, tuple) else v

    return max(matches, key=max_value)


# =========================================
# EXPERIENCE SCORING
# =========================================

# Keep scoring non-binary so the screening layer can still rank partial matches
# when resumes omit an explicit summary but list enough dated experience to infer fit.
def score_years_of_experience(
    extracted_exp: Optional[Union[int, Tuple[int, int]]],
    required_exp: int,
    max_bonus_years: int = 5
) -> int:
    """
    Returns an ATS-style score from 0–100
    """

    if extracted_exp is None:
        # Missing experience should lower confidence without zeroing the score,
        # because many resumes omit a summary while still listing dated roles.
        return 20  # Missing data penalty (not zero)

    # Use the upper bound so "3-5 years" behaves like the strongest stated value.
    years = extracted_exp[1] if isinstance(extracted_exp, tuple) else extracted_exp

    if years >= required_exp:
        bonus = min(years - required_exp, max_bonus_years)
        return min(100, 80 + int((bonus / max_bonus_years) * 20))

    # Partial credit keeps near-matches rankable instead of collapsing them to zero.
    ratio = years / required_exp
    return max(30, int(ratio * 80))
